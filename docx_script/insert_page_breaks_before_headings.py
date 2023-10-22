from docx import Document

# 打开现有的Word文档
doc = Document('demo.docx')  # 将 'your_document.docx' 替换为你的文档路径

# 创建一个新文档用于保存修改后的内容
new_doc = Document()

# 用于标识是否前面是一级标题
previous_was_heading = False

# 遍历文档的段落
for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Heading 1':
        # 打印一级标题内容
        heading_text = paragraph.text
        print(heading_text)

        # 标记前一个段落是一级标题
        previous_was_heading = True

        # 插入分页符到新文档
        if previous_was_heading:
            new_doc.add_page_break()

        # 将一级标题添加到新文档
        new_doc.add_paragraph(heading_text, style='Heading 1')
    else:
        # 复制原文档的段落到新文档
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

        # 标记前一个段落不是一级标题
        previous_was_heading = False

# 保存修改后的文档
new_doc.save('document_with_page_breaks.docx')  # 保存为一个新的文件

print("Page breaks inserted before each first-level heading.")
