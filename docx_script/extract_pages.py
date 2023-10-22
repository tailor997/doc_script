from docx import Document

def save_pages_range(input_path, output_path, start_page, end_page):
    # 打开现有的Word文档
    doc = Document(input_path)

    # 创建一个新文档用于保存指定范围的页
    new_doc = Document()

    # 用于标识当前页码
    current_page = 1

    # 用于标识是否在指定范围内
    in_range = False

    # 用于标识是否已经跳过封面页
    cover_skipped = False

    # 遍历文档的段落
    for paragraph in doc.paragraphs:
        if not cover_skipped:
            # 跳过封面页
            cover_skipped = True
            continue

        if 'PAGE' in paragraph.text:
            current_page += 1

            if current_page >= start_page and current_page <= end_page:
                in_range = True
            else:
                in_range = False
        elif in_range:
            new_paragraph = new_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

    # 保存修改后的文档
    new_doc.save(output_path)

# 指定输入文档、输出文档以及起始和结束页码
input_file = 'PAN1070_GPIO.docx'  # 替换为你的输入文件路径
output_file = 'output_pages.docx'  # 替换为你的输出文件路径
start_page = 1
end_page = 3  # 从第A页到第B页

# 调用函数保存指定范围的页
save_pages_range(input_file, output_file, start_page, end_page)
