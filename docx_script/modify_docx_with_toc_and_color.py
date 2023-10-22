from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor  # 引入 RGBColor

# 打开现有的Word文档
doc = Document('demo.docx')

# 创建一个自定义样式，例如 "TOC_Style"
toc_style = doc.styles.add_style('TOC_Style', WD_STYLE_TYPE.PARAGRAPH)

# 设置该样式的属性，例如字体大小和对齐方式
toc_style.font.size = Pt(12)
toc_style.paragraph_format.alignment = 1  # 居中对齐，1 表示居中

# 添加目录
doc.add_paragraph("Table of Contents", style='TOC_Style')

# 遍历文档的段落，找到一级标题并修改字体颜色为黄色
for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Heading 1':
        for run in paragraph.runs:
            font = run.font
            font.color.rgb = RGBColor(255, 255, 0)  # 设置字体颜色为黄色

# 保存修改后的文档
doc.save('modified_existing_document.docx')

print("Existing document modified with a table of contents and font color for first-level headings.")
