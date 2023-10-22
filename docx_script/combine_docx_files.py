from docx import Document

def combine_docx_files(file1, file2, output_file):
    # 打开第一个docx文件
    doc1 = Document(file1)

    # 打开第二个docx文件
    doc2 = Document(file2)

    # 创建一个新文档用于保存合并的内容
    combined_doc = Document()

    # 复制第一个文档的内容到合并文档
    for element in doc1.element.body:
        combined_doc.element.body.append(element)

    # 复制第二个文档的内容到合并文档
    for element in doc2.element.body:
        combined_doc.element.body.append(element)

    # 保存合并后的文档
    combined_doc.save(output_file)

# 指定两个输入文档和输出文档的文件路径
input_file1 = 'cover template.docx'  # 替换为第一个输入文件路径
input_file2 = 'output_pages.docx'  # 替换为第二个输入文件路径
output_file = 'combined_file.docx'  # 替换为输出文件路径

# 调用函数将两个文档拼接成一个
combine_docx_files(input_file1, input_file2, output_file)
